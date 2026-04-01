"""
DOCX Extractor — Structured extraction from Word documents.

Extracts:
  • Headings (with levels)
  • Paragraphs preserving structure
  • Tables as markdown
  • Images with optional Vision API descriptions
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import List, Optional

from ..models import ExtractedDocument, ExtractedElement, ExtractedPage, ImageReference

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract structured content from .docx files."""

    def __init__(self, *, image_describer=None, image_store=None):
        self.image_describer = image_describer
        self.image_store = image_store

    def extract(self, file_path: str) -> ExtractedDocument:
        """Extract all content from a DOCX file."""
        path = Path(file_path)
        doc = ExtractedDocument(
            source_path=str(path),
            file_type="docx",
            title=path.stem,
            metadata={"filename": path.name},
        )

        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed — run: pip install python-docx")
            return doc

        try:
            docx = DocxDocument(file_path)
            elements: List[ExtractedElement] = []
            page_num = 1  # DOCX doesn't have real pages; we simulate per major heading

            # ── Core properties ──
            if docx.core_properties:
                if docx.core_properties.title:
                    doc.title = docx.core_properties.title
                if docx.core_properties.author:
                    doc.metadata["author"] = docx.core_properties.author

            # ── Paragraphs & headings ──
            for para in docx.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower()

                if "heading" in style_name:
                    # Extract heading level
                    level = 1
                    m = re.search(r"heading\s*(\d)", style_name)
                    if m:
                        level = int(m.group(1))
                    if level == 1:
                        page_num += 1  # Simulate page break on major headings
                    elements.append(ExtractedElement(
                        content=text,
                        element_type="heading",
                        page_number=page_num,
                        metadata={"heading_level": level, "style": para.style.name},
                    ))
                elif "list" in style_name or text.startswith(("•", "-", "●", "*")):
                    elements.append(ExtractedElement(
                        content=text,
                        element_type="list",
                        page_number=page_num,
                    ))
                else:
                    elements.append(ExtractedElement(
                        content=text,
                        element_type="text",
                        page_number=page_num,
                    ))

            # ── Tables ──
            for table in docx.tables:
                md = self._table_to_markdown(table)
                if md:
                    elements.append(ExtractedElement(
                        content=md,
                        element_type="table",
                        page_number=page_num,
                        metadata={"rows": len(table.rows), "cols": len(table.columns)},
                    ))

            # ── Images (Phase 3: always extract, even without image_describer) ──
            self._extract_images(docx, elements, path.name, str(path))

            # Group elements into pages (by page_number)
            pages_map: dict[int, ExtractedPage] = {}
            for elem in elements:
                pn = elem.page_number
                if pn not in pages_map:
                    pages_map[pn] = ExtractedPage(page_number=pn)
                pages_map[pn].elements.append(elem)

            for pn in sorted(pages_map):
                ep = pages_map[pn]
                ep.composite_text = "\n\n".join(e.content for e in ep.elements)
                ep.raw_text = ep.composite_text
                doc.pages.append(ep)

            doc.metadata["total_pages"] = len(doc.pages)
            logger.info(f"✅ Extracted {len(elements)} elements from {path.name}")

        except Exception as exc:
            logger.error(f"DOCX extraction failed for {path.name}: {exc}", exc_info=True)

        return doc

    # ────────────────────────────────────────────────────
    #  Images
    # ────────────────────────────────────────────────────
    def _extract_images(self, docx, elements: list, filename: str, file_path: str):
        """Extract images from DOCX relationships, save to disk, create ImageReferences.

        Phase 3 upgrade: works without image_describer, saves via ImageStore,
        creates proper ImageReference objects with captions from context.
        """
        try:
            # Gather headings from already-extracted elements for context
            heading_texts = [e.content for e in elements if e.element_type == "heading"]
            context_summary = " > ".join(heading_texts[-3:]) if heading_texts else filename

            img_idx = 0
            for rel in docx.part.rels.values():
                if "image" not in rel.reltype:
                    continue
                try:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    if len(img_bytes) < 2000:  # Skip tiny icons/bullets
                        continue

                    # Determine image dimensions
                    img_width, img_height = self._get_image_dimensions(img_bytes)
                    if img_width < 100 or img_height < 100:
                        continue

                    page_num = 1  # DOCX doesn't have real page numbers

                    # Save image to disk via ImageStore
                    image_path = ""
                    if self.image_store:
                        image_path = self.image_store.save(
                            img_bytes, file_path, page_num, img_idx,
                        )

                    # Vision API description (optional)
                    description = ""
                    if self.image_describer:
                        description = self.image_describer.describe(
                            img_bytes, context=f"Document: {filename}"
                        )

                    # Build contextual caption
                    caption = f"{context_summary} | Image {img_idx + 1} from {filename}"

                    # Build ImageReference
                    img_ref = ImageReference(
                        image_path=image_path,
                        page_number=page_num,
                        caption=caption,
                        ocr_text="",
                        description=description,
                        width=img_width,
                        height=img_height,
                    )

                    display_text = description or caption
                    elements.append(ExtractedElement(
                        content=f"[Image: {display_text}]",
                        element_type="image",
                        page_number=page_num,
                        metadata={
                            "image_index": img_idx,
                            "width": img_width,
                            "height": img_height,
                            "description": description,
                            "caption": caption,
                            "image_path": image_path,
                            "image_ref": img_ref.to_dict(),
                        },
                    ))
                    img_idx += 1
                except Exception as exc:
                    logger.debug(f"Skipping DOCX image {img_idx}: {exc}")
        except Exception as exc:
            logger.debug(f"DOCX image extraction error: {exc}")

    @staticmethod
    def _get_image_dimensions(img_bytes: bytes) -> tuple:
        """Get image width and height from bytes."""
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            return img.width, img.height
        except Exception:
            # If PIL not available, return reasonable defaults
            return 200, 200

    # ────────────────────────────────────────────────────
    #  Table → Markdown
    # ────────────────────────────────────────────────────
    @staticmethod
    def _table_to_markdown(table) -> str:
        """Convert a python-docx Table to Markdown."""
        rows_data = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            rows_data.append(cells)
        if not rows_data:
            return ""

        header = rows_data[0]
        md = "| " + " | ".join(header) + " |\n"
        md += "| " + " | ".join("---" for _ in header) + " |\n"
        for row in rows_data[1:]:
            padded = row + [""] * (len(header) - len(row)) if len(row) < len(header) else row[:len(header)]
            md += "| " + " | ".join(padded) + " |\n"
        return md.strip()
